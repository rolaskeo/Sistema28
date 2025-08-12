from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from datetime import datetime

def crear_word(titulo_consulta, cuerpo_secciones, referencias, ruta_salida):
    """
    Genera un documento Word con formato PTN adaptado a Sistema 28.

    Args:
        titulo_consulta (str): Tema de consulta.
        cuerpo_secciones (list): Lista de secciones con formato:
                                 [("I. ANTECEDENTES", ["texto1", "texto2"]), ...]
                                 donde cada párrafo puede contener referencias tipo [1].
        referencias (dict): Diccionario {numero: texto} con notas al final.
        ruta_salida (str): Ruta donde guardar el .docx

    Returns:
        str: Ruta final del archivo generado.
    """

    doc = Document()

    # Encabezado
    encabezado = doc.add_paragraph("S28 v.1.0")
    encabezado.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = encabezado.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    # Título principal
    titulo = doc.add_paragraph("DICTAMEN JURÍDICO")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    # Fecha y tema de consulta
    fecha = doc.add_paragraph(datetime.now().strftime("%d/%m/%Y"))
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha.runs[0].font.name = "Times New Roman"
    fecha.runs[0].font.size = Pt(10)

    tema = doc.add_paragraph(f"Tema de consulta: {titulo_consulta}")
    tema.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tema.runs[0].font.name = "Times New Roman"
    tema.runs[0].font.size = Pt(12)

    # Secciones y numeración de párrafos
    for sec_idx, (titulo_sec, parrafos) in enumerate(cuerpo_secciones, start=1):
        # Título de sección
        sec = doc.add_paragraph(titulo_sec.upper())
        sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sec_run = sec.runs[0]
        sec_run.font.name = "Times New Roman"
        sec_run.font.size = Pt(12)
        sec_run.bold = True

        # Párrafos numerados
        for par_idx, texto in enumerate(parrafos, start=1):
            par = doc.add_paragraph()
            par.paragraph_format.first_line_indent = Cm(1.25)
            par.paragraph_format.line_spacing = 1.5
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = par.add_run(f"{sec_idx}.{par_idx} {texto}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # Firma
    doc.add_paragraph("")
    firma = doc.add_paragraph("Dr. Rolando Keumurdji Rizzuti")
    firma.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_firma = firma.runs[0]
    run_firma.font.name = "Times New Roman"
    run_firma.font.size = Pt(12)

    # Media línea separadora
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Notas al final
    if referencias:
        notas = doc.add_paragraph("Referencias:")
        notas.alignment = WD_ALIGN_PARAGRAPH.LEFT
        notas.runs[0].bold = True
        notas.runs[0].font.name = "Times New Roman"
        notas.runs[0].font.size = Pt(10)

        for num, ref in referencias.items():
            ref_par = doc.add_paragraph(f"[{num}] {ref}")
            ref_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ref_par.paragraph_format.first_line_indent = Cm(0)
            ref_par.runs[0].font.name = "Times New Roman"
            ref_par.runs[0].font.size = Pt(10)

    doc.save(ruta_salida)
    return ruta_salida
