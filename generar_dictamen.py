#!/usr/bin/env python3
# generar_dictamen.py — versión consolidada Sistema 28
# - Recuperación FAISS + embeddings
# - Llamada remota OpenRouter/DeepSeek (usa API_KEY en .env como "API_KEY")
# - Exportación a .docx con formato PTN (titulos centrados, numeración 1.1, notas al final tras firma)
# - Guardado en SQLite (dictamenes, antecedentes_relevantes, documentos_subidos)
# - CLI: generar dictamen / solo antecedentes / subir documento

import os
import json
import time
import sqlite3
import requests
import re
from datetime import datetime
from dotenv import load_dotenv

import numpy as np
import faiss
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pathlib import Path

# --- v3: búsqueda híbrida desde s28 ---
import sys
sys.path.append(os.getcwd())  # para que VSCode/Streamlit encuentre el paquete local 's28'
from s28.rerank import search as v3_search

# ---------------- CONFIG ----------------
load_dotenv()

API_KEY = os.getenv("API_KEY")  # <- tu .env debe contener: API_KEY=xxxxx
MODEL_ID = os.getenv("MODEL_ID", "deepseek/deepseek-chat")
OPENROUTER_API_URL = os.getenv("OPENROUTER_API_URL", "https://openrouter.ai/api/v1/chat/completions")

FAISS_INDEX_PATH = os.getenv("FAISS_INDEX_PATH", "faiss_index.bin")
FAISS_METADATA_PATH = os.getenv("FAISS_METADATA_PATH", "faiss_metadata.json")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")

DB_PATH = os.getenv("DATABASE_PATH", "sistema28.db")
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "dictamenes")
TEMPLATE_PATH = os.getenv("DICTAMEN_TEMPLATE", "plantillas/dictamen_ptn.docx")
DICTAMEN_ESTRUCTURA_PATH = os.getenv("DICTAMEN_ESTRUCTURA", "plantilla/estructura.json")
MAX_CONTEXT_CHUNKS = int(os.getenv("MAX_CONTEXT_CHUNKS", 5))
SIMILARITY_THRESHOLD = float(os.getenv("SIMILARITY_THRESHOLD", 0.35))

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------- helpers docx ----------------
def add_page_number_field(paragraph, prefix_text="Página "):
    run = paragraph.add_run(prefix_text)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)

def normalize_text(s: str) -> str:
    return re.sub(r'\s+', ' ', (s or "")).strip()

# ---------------- load embeddings & faiss ----------------
print(f"📥 Cargando modelo de embeddings: {EMBEDDING_MODEL} ...")
embed_model = SentenceTransformer(EMBEDDING_MODEL)

if not os.path.exists(FAISS_INDEX_PATH) or not os.path.exists(FAISS_METADATA_PATH):
    raise FileNotFoundError(f"Faltan archivos FAISS: {FAISS_INDEX_PATH} y/o {FAISS_METADATA_PATH}")

index = faiss.read_index(FAISS_INDEX_PATH)
with open(FAISS_METADATA_PATH, "r", encoding="utf-8") as f:
    raw_meta = json.load(f)

# Normalize metadata to dict[int] -> meta
metadata = {}
if isinstance(raw_meta, dict):
    for k, v in raw_meta.items():
        try:
            metadata[int(k)] = v
        except:
            metadata[k] = v
elif isinstance(raw_meta, list):
    for i, item in enumerate(raw_meta):
        metadata[i] = item
else:
    raise RuntimeError("Formato inesperado en faiss_metadata.json")

# ---------------- DB init ----------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS dictamenes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        consulta TEXT,
        ruta_docx TEXT,
        fecha TEXT,
        modelo TEXT,
        top_chunks JSON
    )""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS antecedentes_relevantes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        dictamen_id INTEGER,
        chunk_idx INTEGER,
        document_path TEXT,
        tipo_documento TEXT,
        organismo_emisor TEXT,
        fecha_documento TEXT,
        snippet TEXT,
        similitud REAL,
        valoracion INTEGER,
        comentario TEXT,
        FOREIGN KEY(dictamen_id) REFERENCES dictamenes(id)
    )""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS documentos_subidos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        titulo TEXT,
        tipo_documento TEXT,
        palabras_clave TEXT,
        relevancia INTEGER,
        contenido TEXT,
        fecha_subida TEXT
    )""")
    conn.commit()
    conn.close()

init_db()

# ---------------- FAISS search ----------------
def get_topk_chunks(query_text, k=MAX_CONTEXT_CHUNKS):
    q_emb = embed_model.encode([query_text], convert_to_numpy=True)
    q_vec = np.array(q_emb, dtype="float32")
    D, I = index.search(q_vec, k)
    hits = []
    for dist_arr, idx_arr in zip(D, I):
        for dist, idx in zip(dist_arr, idx_arr):
            if idx == -1:
                continue
            meta = metadata.get(int(idx)) if int(idx) in metadata else metadata.get(str(idx))
            if not meta:
                continue
            # If index type is L2 distance, smaller is better. We'll expose raw distance and later use semantic sim if needed.
            hits.append({
                "idx": int(idx),
                "similitud": float(dist),
                "document_path": meta.get("document_path", ""),
                "tipo_documento": meta.get("tipo_documento", meta.get("document_type", "")),
                "organismo_emisor": meta.get("organismo_emisor", ""),
                "fecha_documento": meta.get("fecha_documento", ""),
                "chunk_text": normalize_text(meta.get("chunk_text", ""))
            })
    return hits

# --- NUEVO (v3): recuperar top-k con FTS→coseno→meta→MMR ---
def get_topk_chunks_v3(db_path: str, query: str, k: int = MAX_CONTEXT_CHUNKS, candidates: int = 200):
    conn = sqlite3.connect(db_path)
    conn.row_factory = lambda c,r:{d[0]: r[i] for i,d in enumerate(c.description)}
    rows = v3_search(conn, query, topk=k, candidates=candidates, with_snippet=True)

    # Enriquecer con issuer/year en un solo query (para encabezados)
    doc_ids = list({int(r["doc_id"]) for r in rows})
    meta = {}
    if doc_ids:
        qmarks = ",".join("?" for _ in doc_ids)
        mrows = conn.execute(f"SELECT id, issuer, year FROM documents_v3 WHERE id IN ({qmarks})", doc_ids).fetchall()
        meta = {int(r["id"]): {"issuer": r["issuer"], "year": r["year"]} for r in mrows}

    conn.close()

    hits = []
    for r in rows:
        dm = meta.get(int(r["doc_id"]), {})
        issuer = dm.get("issuer") or ""
        year   = dm.get("year") or ""
        hits.append({
            # campos compatibles con tu UI/flujo
            "id":       r["chunk_id"],     # id del chunk para citas internas
            "doc_id":   r["doc_id"],
            "titulo":   r["title"],
            "page":     r["page"],         # en PDF viene número; DOCX suele ser None
            "chunk_text": r.get("text") or r.get("snippet") or "",
            "preview":  r.get("snippet") or "",
            "score":    r["score"],
            "path":     r["path"],
            # mapeo a los nombres legacy (para encabezados existentes)
            "organismo_emisor": issuer,
            "fecha_documento":  str(year) if year else "",
            "tipo_documento":   "",        # no lo tenemos en v3; queda vacío
        })
    return hits

# ---------------- remote model call (OpenRouter / DeepSeek) ----------------
def call_remote_model(system_msg, user_msg, max_tokens=1024, temperature=0.2, attempts=3, timeout=30):
    if not API_KEY:
        raise RuntimeError("API_KEY no encontrada en variables de entorno (.env).")
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": MODEL_ID,
        "messages": [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg}
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False
    }
    last_err = None
    for attempt in range(1, attempts + 1):
        try:
            r = requests.post(OPENROUTER_API_URL, headers=headers, json=payload, timeout=timeout)
            r.raise_for_status()
            j = r.json()
            if "choices" in j and len(j["choices"]) > 0:
                ch = j["choices"][0]
                if "message" in ch and isinstance(ch["message"], dict):
                    return ch["message"].get("content", "")
                if "text" in ch:
                    return ch["text"]
            return json.dumps(j, ensure_ascii=False)
        except Exception as e:
            print(f"[warning] intento {attempt} falló: {e}. reintentando en {2**(attempt-1)}s...")
            last_err = e
            time.sleep(2**(attempt-1))
    raise RuntimeError(f"No se obtuvo respuesta de la API remota. Último error: {last_err}")

# ---------------- associate paragraphs with sources (semantic) ----------------
def associate_paragraphs_with_sources(full_text, top_k_per_par=3, similarity_threshold=SIMILARITY_THRESHOLD):
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', full_text) if p.strip()]
    results = []
    chunk_emb_cache = {}
    for p in paragraphs:
        p_emb = embed_model.encode([p], convert_to_numpy=True)[0]
        D, I = index.search(np.array([p_emb], dtype="float32"), top_k_per_par)
        best_meta = None
        best_sim = -1.0
        # compute cosine sims against cached embeddings (more robust)
        for idx in I[0]:
            if idx == -1:
                continue
            meta = metadata.get(int(idx))
            if not meta:
                continue
            chunk_text = normalize_text(meta.get("chunk_text", ""))
            if idx not in chunk_emb_cache:
                chunk_emb_cache[idx] = embed_model.encode([chunk_text], convert_to_numpy=True)[0]
            sim = util.cos_sim(p_emb, chunk_emb_cache[idx]).item()
            if sim > best_sim:
                best_sim = sim
                best_meta = meta
        results.append((p, best_meta if best_sim >= similarity_threshold else None, float(best_sim)))
    return results

# --- NUEVO: completar referencias vacías con hits v3 ---
def backfill_refs_from_v3(associated_paragraphs, v3_hits, sim_threshold=0.30, prefer_v3=True, margin=0.05):
    """
    Si 'prefer_v3' es True, V3 puede REEMPLAZAR la meta encontrada por FAISS cuando:
      - no hay meta FAISS, o
      - el mejor match de V3 tiene similitud >= sim_threshold y es >= (sim_FAISS + margin).
    """
    if not v3_hits:
        return associated_paragraphs

    v3_texts, v3_meta = [], []
    for h in v3_hits:
        txt = (h.get("chunk_text") or h.get("preview") or "").strip()
        v3_texts.append(txt)
        v3_meta.append({
            "document_path": h.get("path") or h.get("document_path") or "",
            "tipo_documento": h.get("tipo_documento") or "",
            "organismo_emisor": h.get("organismo_emisor") or "",
            "fecha_documento": h.get("fecha_documento") or "",
            "chunk_text": txt,
        })
    if not any(v3_texts):
        return associated_paragraphs

    v3_vecs = embed_model.encode(v3_texts, convert_to_numpy=True)

    out = []
    for (p, meta, sim0) in associated_paragraphs:
        p_vec = embed_model.encode([p], convert_to_numpy=True)[0]
        sims = v3_vecs @ p_vec
        best = int(np.argmax(sims))
        best_score = float(sims[best])

        # ¿reemplazar?
        do_override = (not meta or not meta.get("document_path")) or (prefer_v3 and best_score >= max(sim0 or 0.0, 0.0) + margin)
        if best_score >= sim_threshold and do_override:
            out.append((p, v3_meta[best], best_score))
        else:
            out.append((p, meta, sim0))
    return out

# ---------------- create docx PTN (strict format) ----------------
ROMAN_RE = re.compile(r'^\s*([IVXLCDM]+)\.\s*(.*)', flags=re.IGNORECASE)

def create_docx_ptn(consulta, associated_paragraphs):
    """
    - Titles (I. ANTECEDENTES, II. ANÁLISIS, III. CONCLUSIÓN) are centered, uppercase, roman.
    - Paragraph numbering: section.paragraph (1.1, 1.2, 2.1...)
    - First line indent 1.25 cm. Justified text. Times New Roman 12. Title font 14.
    - After signature: horizontal separator and 'Notas y Referencias' (size 10).
    - References are collected from associated_paragraphs: unique by (path,tipo,org,fecha).
    """
    # Intentar abrir la plantilla definida en .env; si falla, fallback a documento en blanco.
    use_template = False
    try:
        doc = Document(TEMPLATE_PATH)
        use_template = True
    except Exception:
        doc = Document()
    # >>> NUEVO: solo si NO hay plantilla, aplicamos formato PTN por defecto
    if not use_template:
        sec = doc.sections[0]
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Portada
        title_p = doc.add_paragraph()
        title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title_p.add_run("Dictamen Jurídico")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        doc.add_paragraph("")
        fecha_p = doc.add_paragraph()
        fecha_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fecha_p.add_run(datetime.now().strftime("%d de %B de %Y")).font.size = Pt(12)

        tema = consulta.strip() if len(consulta) < 250 else (consulta[:247] + "...")
        tema_p = doc.add_paragraph()
        tema_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tema_p.add_run(tema).font.size = Pt(12)

        doc.add_page_break()

        # Header & footer
        for s in doc.sections:
            if s.header.paragraphs:
                s.header.paragraphs[0].text = "S28 v.3.3"
                s.header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if s.footer.paragraphs:
                f = s.footer.paragraphs[0]
                f.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_page_number_field(f, prefix_text="Página ")
# <<< FIN NUEVO

    # Body with section detection and numbering
    used_refs = []  # list of keys (path,tipo,org,fecha)
    section_counter = 0
    paragraph_counter = 0

    for block, meta, score in associated_paragraphs:
        block = block.strip()
        # detect section header like "I. Antecedentes" or "I. ANTECEDENTES"
        m = ROMAN_RE.match(block)
        if m:
            section_counter += 1
            paragraph_counter = 0
            title_text = m.group(0).strip()
            # ensure uppercase title after roman numeral
            # Create centered bold uppercase title like "I. ANTECEDENTES"
            sec_par = doc.add_paragraph()
            sec_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = sec_par.add_run(title_text.upper())
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # else regular paragraph
        paragraph_counter += 1
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(6)
        pf.first_line_indent = Cm(1.25)

        # paragraph number "section.paragraph "
        num_run = p.add_run(f"{section_counter}.{paragraph_counter} ")
        num_run.bold = True
        num_run.font.name = "Times New Roman"
        num_run.font.size = Pt(12)

        # write paragraph text (clean ** markers)
        text_clean = normalize_text(block.replace("**", ""))
        p.add_run(text_clean).font.name = "Times New Roman"

        # handle references: if meta exists attach a superscript consecutive number.
        if meta:
            key = (meta.get("document_path",""), meta.get("tipo_documento", meta.get("document_type","")), meta.get("organismo_emisor",""), meta.get("fecha_documento",""))
            if key in used_refs:
                ref_num = used_refs.index(key) + 1
            else:
                used_refs.append(key)
                ref_num = len(used_refs)
            sup = p.add_run(f"{ref_num}")
            sup.font.superscript = True
            sup.font.size = Pt(10)

    # Signature
    doc.add_paragraph("\n")
    sign_par = doc.add_paragraph()
    sign_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    srun = sign_par.add_run("Dr. Rolando Keumurdji Rizzuti")
    srun.bold = True
    srun.font.size = Pt(12)

    # Separator line then notes AFTER signature (user requested notes after signature)
    sep = doc.add_paragraph()
    sep.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sep_run = sep.add_run("-" * 60)
    sep_run.font.size = Pt(8)

    # Notas y Referencias
    doc.add_page_break()
    doc.add_heading("Notas y Referencias", level=1)
    last_path = None
    for i, key in enumerate(used_refs, start=1):
        path, tipo, org, fecha = key
        ref_par = doc.add_paragraph()
        # if same path as previous -> Ibíd.
        if last_path and path == last_path:
            r = ref_par.add_run(f"[{i}] Ibíd.")
            r.bold = True
            r.font.size = Pt(10)
        else:
            # --- fallbacks elegantes si faltan metadatos ---
            tipo  = tipo or "Documento"
            if not org:
                org = Path(path).parent.name if path else ""
            fecha = str(fecha) if fecha else ""
            if not org and path:
                org = Path(path).stem[:80]

            # encabezado de la referencia
            b = ref_par.add_run(f"[{i}] {tipo} - {org} - {fecha}\n")
            b.bold = True
            b.font.size = Pt(10)

            # mostrar solo el nombre de archivo (más legible que el path completo)
            ref_par.add_run(f"Fuente: {Path(path).name if path else '—'}\n").font.size = Pt(10)

            # snippet (robusto ante 'found=None')
            snippet = ""
            if path:
                try:
                    found = next((m for m in metadata.values() if m.get("document_path", "") == path), None)
                    if found:
                        ct = normalize_text(found.get("chunk_text", ""))
                        if ct:
                            snippet = ct[:350] + ("..." if len(ct) > 350 else "")
                except Exception:
                    pass
            if snippet:
                ref_par.add_run(snippet).font.size = Pt(10)
        last_path = path

    # save
    filename = os.path.join(OUTPUT_DIR, f"dictamen_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
    doc.save(filename)
    return filename

# ---------------- DB save ----------------
def save_dictamen_record(consulta, ruta_docx, modelo, top_hits):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO dictamenes (consulta, ruta_docx, fecha, modelo, top_chunks) VALUES (?, ?, ?, ?, ?)",
        (consulta, ruta_docx, datetime.now().isoformat(), modelo, json.dumps(top_hits, ensure_ascii=False))
    )
    dictamen_id = cur.lastrowid
    for hit in top_hits:
        cur.execute("""
            INSERT INTO antecedentes_relevantes
            (dictamen_id, chunk_idx, document_path, tipo_documento, organismo_emisor, fecha_documento, snippet, similitud)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, 
            (
                dictamen_id,
                hit.get("idx") or hit.get("id") or hit.get("chunk_id"),
                hit.get("document_path") or hit.get("path"),
                hit.get("tipo_documento"),
                hit.get("organismo_emisor"),
                hit.get("fecha_documento"),
                (hit.get("chunk_text")[:800] if hit.get("chunk_text") else ""),
                hit.get("similitud", hit.get("score")),
            )
    )
    conn.commit()
    conn.close()
    return dictamen_id

# ---------------- valorar antecedentes (CLI) ----------------
def cli_valorar_antecedentes(dictamen_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, chunk_idx, document_path, snippet, similitud, valoracion FROM antecedentes_relevantes WHERE dictamen_id = ?", (dictamen_id,))
    rows = cur.fetchall()
    if not rows:
        print("No hay antecedentes asociados a ese dictamen.")
        conn.close()
        return
    for r in rows:
        aid, chunk_idx, path, snippet, sim, val = r
        print(f"\nID: {aid} | chunk_idx: {chunk_idx} | similitud: {sim:.4f}\n{(snippet or '')[:300]}...\nValor actual: {val}")
        resp = input("¿Valorás este antecedente? (1-5 o Enter para saltar): ").strip()
        if resp and resp.isdigit() and 1 <= int(resp) <= 5:
            cur.execute("UPDATE antecedentes_relevantes SET valoracion = ? WHERE id = ?", (int(resp), aid))
    conn.commit()
    conn.close()
    print("Valoraciones guardadas.")

# ---------------- subir documento manual (interfaz) ----------------
def subir_documento_manual_cli():
    print("Subir documento - registro manual (no actualiza FAISS automáticamente).")
    titulo = input("Título (o nombre de archivo): ").strip()
    tipo = input("Tipo de documento (sentencia, norma, dictamen, doctrina, opinión propia, informe técnico): ").strip()
    kws = input("Palabras clave (3-5) separadas por coma. Use '_' para espacios, ej: codigo_civil_y_comercial: ").strip()
    relev = input("Relevancia (1-5): ").strip()
    contenido = input("Contenido / texto (pegar o corto resumen):\n")
    if not (titulo and tipo and kws and contenido):
        print("Faltan datos obligatorios. Abortando.")
        return
    if len([k for k in kws.split(",") if k.strip()]) < 3:
        print("Debe ingresar entre 3 y 5 palabras clave. Abortando.")
        return
    try:
        relev_int = int(relev)
        if not (1 <= relev_int <= 5):
            raise ValueError()
    except:
        relev_int = 3
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO documentos_subidos (titulo, tipo_documento, palabras_clave, relevancia, contenido, fecha_subida)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (titulo, tipo, kws, relev_int, contenido, datetime.now().isoformat()))
    conn.commit()
    conn.close()
    print("Documento registrado en 'documentos_subidos'.")
    print("IMPORTANTE: Para que este documento influya en búsquedas FAISS debe indexarse (ejecutar proceso de chunking + reindexación).")

# ---------------- CLI principal ----------------
def option_generate_dictamen():
    consulta = input("📌 Ingresá tu consulta jurídica:\n>>> ").strip()
    if not consulta:
        print("Consulta vacía. Abortando.")
        return
    print(f"📚 Recuperando hasta {MAX_CONTEXT_CHUNKS} bloques relevantes...")
    hits = get_topk_chunks_v3(DB_PATH, consulta, k=MAX_CONTEXT_CHUNKS)
    if not hits:
        print("No se encontraron antecedentes relevantes.")
        return

    # build context
    context_parts = []
    for i, hit in enumerate(hits, start=1):
        snippet = hit.get("chunk_text","")[:1200]
        header = f"[{i}] {hit.get('tipo_documento','')} | {hit.get('organismo_emisor','')} | {hit.get('fecha_documento','')}\n"
        context_parts.append(header + snippet)
    contexto_text = "\n\n".join(context_parts)

    # Cargar estructura del dictamen desde JSON (si existe)
    try:
        with open(DICTAMEN_ESTRUCTURA_PATH, "r", encoding="utf-8") as f:
            esquema = (json.load(f).get("secciones") or [])
            esquema = [s for s in esquema if isinstance(s, str) and s.strip()]
    except Exception:
        esquema = []

    if not esquema:
        esquema = ["I. Antecedentes", "II. Análisis", "III. Conclusión"]  # fallback breve

    epigrafes = "\n- " + "\n- ".join(esquema)  # para mostrarle al LLM “exactamente” estos títulos

    system_prompt = (
        "Sos un abogado experto en derecho administrativo argentino. "
        "Utilizá únicamente el contexto que te brindo para redactar un proyecto de dictamen jurídico. "
        "No inventes normas ni citas: si una norma es mencionada, debe estar textual en el contexto provisto. "
        "Estructurá el dictamen exactamente con estos epígrafes:" + epigrafes + " "
        "Usá citas en formato [n] donde n refiere a las fuentes recuperadas. "
        "Aplicá el Manual de Estilo PTN 2023 en lo que corresponda."
    )

    user_prompt = (
        f"Consulta: {consulta}\n\n"
        f"Contexto relevante:\n{contexto_text}\n\n"
        "Redactá el dictamen técnico-jurídico."
    )

    print("🧠 Solicitando redacción al modelo (remoto)...")
    try:
        response_text = call_remote_model(system_prompt, user_prompt, max_tokens=1200, temperature=0.2, attempts=3, timeout=60)
    except Exception as e:
        print("Error al llamar al modelo:", e)
        return

    print("\n📝 Respuesta (preview):\n")
    print(response_text[:2000] + ("..." if len(response_text) > 2000 else ""))

    print("🔎 Asociando párrafos con fuentes (semántico)...")
    associated = associate_paragraphs_with_sources(response_text, top_k_per_par=3, similarity_threshold=SIMILARITY_THRESHOLD)
    
    # NUEVO: completar referencias con v3
    associated = backfill_refs_from_v3(associated, hits)

    print("💾 Generando documento Word (formato PTN)...")
    filename = create_docx_ptn(consulta, associated)
    print("Documento generado:", filename)

    dictamen_id = save_dictamen_record(consulta, filename, MODEL_ID, hits)
    print("Dictamen guardado en la base con id:", dictamen_id)

    ans = input("¿Querés valorar los antecedentes ahora? (s/n): ").strip().lower()
    if ans == "s":
        cli_valorar_antecedentes(dictamen_id)

def option_antecedentes_only():
    consulta = input("📌 Ingresá tu búsqueda de antecedentes (consulta libre):\n>>> ").strip()
    if not consulta:
        print("Consulta vacía. Abortando.")
        return
    hits = get_topk_chunks_v3(DB_PATH, consulta, k=MAX_CONTEXT_CHUNKS)
    if not hits:
        print("No se encontraron antecedentes.")
        return
    print(f"\n📚 Se encontraron {len(hits)} antecedentes (resumen):\n")
    for i, h in enumerate(hits, start=1):
        print(f"---[{i}]---")
        print(f"{h.get('tipo_documento')} | {h.get('organismo_emisor')} | {h.get('fecha_documento')}")
        print(h.get('chunk_text')[:800] + ("..." if len(h.get('chunk_text',''))>800 else ""))
        val = input("¿Valorás este antecedente? (1-5 o Enter para saltar): ").strip()
        if val.isdigit() and 1 <= int(val) <= 5:
            conn = sqlite3.connect(DB_PATH)
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO antecedentes_relevantes
                (dictamen_id, chunk_idx, document_path, tipo_documento, organismo_emisor, fecha_documento, snippet, similitud, valoracion)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, 
                (
                    None,
                    h.get("idx") or h.get("id") or h.get("chunk_id"),
                    h.get("document_path") or h.get("path"),
                    h.get("tipo_documento"),
                    h.get("organismo_emisor"),
                    h.get("fecha_documento"),
                    (h.get("chunk_text")[:800] if h.get("chunk_text") else ""),
                    h.get("similitud", h.get("score")),
                    int(val),
                ))
            conn.commit()
            conn.close()
    print("Operación finalizada. Las valoraciones guardadas.")

def option_subir_documento():
    subir_documento_manual_cli()

def main_cli():
    print("=== Sistema 28 — Generador de Dictámenes (CLI) ===")
    print("Opciones:")
    print("1) Generar dictamen (Word + guardar en DB)")
    print("2) Solo antecedentes relevantes (ver y valorar)")
    print("3) Subir documento (registro manual en DB)")
    print("0) Salir")
    opt = input("Elegí una opción: ").strip()
    if opt == "1":
        option_generate_dictamen()
    elif opt == "2":
        option_antecedentes_only()
    elif opt == "3":
        option_subir_documento()
    else:
        print("Saliendo.")

if __name__ == "__main__":
    main_cli()
